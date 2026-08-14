#!/usr/bin/env python3
"""Convertit un dossier Markdown en document Word charte Digital·Humans.

Pourquoi ce script : dossier.py genere les documents du comite depuis l'etat en
base — il ne sait pas prendre un Markdown quelconque. Les dossiers rediges par
les directeurs (revue externe du Financier, notes juridiques, etudes) sortent en
Markdown et doivent partir a des destinataires externes. Un .md ne se transmet
pas a une direction financiere.

La charte n'est PAS redefinie ici : elle est importee de dossier.py, qui reste la
source unique (encre, laiton, Georgia pour les titres, Calibri pour le corps).
Une seule surface a maintenir.

Usage : md2dossier.py <fichier.md> [sortie.docx]
        Les images en chemin relatif sont resolues depuis le dossier du .md.
"""
import os, re, sys
sys.path.insert(0, "/workspace/bin")
import dossier as D                      # charte + helpers add/head/img/tbl
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

def filet(doc, couleur=D.BRASS, avant=6, apres=10):
    """Trait horizontal — bordure basse de paragraphe (jamais un tableau)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(avant); p.paragraph_format.space_after = Pt(apres)
    pPr = p._p.get_or_add_pPr(); bd = OxmlElement("w:pBdr"); bt = OxmlElement("w:bottom")
    bt.set(qn("w:val"), "single"); bt.set(qn("w:sz"), "6"); bt.set(qn("w:space"), "1")
    bt.set(qn("w:color"), "%02X%02X%02X" % couleur)
    bd.append(bt); pPr.append(bd)

def encadre(doc, lignes):
    """Bloc cite (>) — fond sable, filet laiton a gauche, italique."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for i, l in enumerate(lignes):
        p = D.add(doc, l, size=10, italic=True, color=D.GREY, space=2)
        p.paragraph_format.left_indent = Cm(0.6)
        pPr = p._p.get_or_add_pPr(); bd = OxmlElement("w:pBdr"); lf = OxmlElement("w:left")
        lf.set(qn("w:val"), "single"); lf.set(qn("w:sz"), "12"); lf.set(qn("w:space"), "8")
        lf.set(qn("w:color"), "%02X%02X%02X" % D.BRASS)
        bd.append(lf); pPr.append(bd)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def puce(doc, txt, num=None):
    p = D.add(doc, ("%s  " % num if num else "\u2022  ") + txt, size=10.5, space=3)
    p.paragraph_format.left_indent = Cm(0.7); p.paragraph_format.first_line_indent = Cm(-0.7)

def align_tableau(sep):
    out = []
    for c in sep:
        c = c.strip()
        out.append(WD_ALIGN_PARAGRAPH.RIGHT if c.endswith(":") and not c.startswith(":")
                   else (WD_ALIGN_PARAGRAPH.CENTER if c.startswith(":") and c.endswith(":")
                         else WD_ALIGN_PARAGRAPH.LEFT))
    return out

def cellules(l):
    l = l.strip()
    if l.startswith("|"): l = l[1:]
    if l.endswith("|"): l = l[:-1]
    return [c.strip() for c in l.split("|")]

def inline(t):
    """Le gras ** est traite par D.add ; on neutralise le reste du balisage."""
    t = re.sub(r"`([^`]+)`", r"\1", t)
    t = re.sub(r"(?<!\*)\*([^*\n]+)\*(?!\*)", r"\1", t)
    t = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", t)
    return t

def convertir(src, dst=None):
    base = os.path.dirname(os.path.abspath(src))
    lignes = open(src, encoding="utf-8").read().split("\n")
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.0)
        s.left_margin = s.right_margin = Cm(2.2)
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(10.5)

    i, premier_titre = 0, True
    para = []
    def vider():
        nonlocal para
        if para:
            D.add(doc, inline(" ".join(para)), size=10.5, space=8)
            para = []
    while i < len(lignes):
        l = lignes[i].rstrip()
        # image + legende
        m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", l.strip())
        if m:
            vider()
            p = m.group(2)
            if not os.path.isabs(p): p = os.path.join(base, p)
            D.img(doc, p, width=15.8)
            if m.group(1):
                D.add(doc, m.group(1), size=8.5, italic=True, color=D.GREY, space=12,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1; continue
        # tableau
        if l.startswith("|") and i + 1 < len(lignes) and re.match(r"^\|[\s:|-]+\|?$", lignes[i+1].strip()):
            vider()
            entetes = cellules(l); alg = align_tableau(cellules(lignes[i+1]))
            i += 2; corps = []
            while i < len(lignes) and lignes[i].strip().startswith("|"):
                corps.append([inline(c) for c in cellules(lignes[i])]); i += 1
            t = D.tbl(doc, [inline(h) for h in entetes], corps)
            for r in t.rows:
                for j, c in enumerate(r.cells):
                    if j < len(alg):
                        for pp in c.paragraphs: pp.alignment = alg[j]
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue
        # titres
        m = re.match(r"^(#{1,3})\s+(.*)", l)
        if m:
            vider()
            niv, txt = len(m.group(1)), inline(m.group(2))
            if niv == 1 and premier_titre:
                premier_titre = False
                D.add(doc, "DIGITAL\u00b7HUMANS", 11, color=D.BRASS, space=2,
                      align=WD_ALIGN_PARAGRAPH.CENTER, font="Georgia")
                filet(doc, apres=14)
                D.add(doc, txt, 20, space=6, align=WD_ALIGN_PARAGRAPH.CENTER, font="Georgia")
                filet(doc, apres=16)
            else:
                if niv == 2: filet(doc, couleur=(0xD8, 0xD2, 0xC6), avant=12, apres=2)
                D.head(doc, txt, niv)
            i += 1; continue
        # bloc cite
        if l.startswith(">"):
            vider(); bloc = []
            courant = []
            while i < len(lignes) and lignes[i].startswith(">"):
                t = inline(lignes[i].lstrip(">").strip())
                if t: courant.append(t)
                elif courant: bloc.append(" ".join(courant)); courant = []
                i += 1
            if courant: bloc.append(" ".join(courant))
            encadre(doc, bloc); continue
        # separateur
        if re.match(r"^-{3,}$", l.strip()):
            vider(); filet(doc, couleur=(0xD8, 0xD2, 0xC6), avant=8, apres=8); i += 1; continue
        # listes
        m = re.match(r"^(\d+)\.\s+(.*)", l) or re.match(r"^([-*])\s+(.*)", l)
        if m:
            vider()
            marque = m.group(1) + "." if m.group(1).isdigit() else None
            corps = [m.group(2)]; i += 1
            while i < len(lignes) and re.match(r"^\s+\S", lignes[i]) \
                  and not re.match(r"^\s*(\d+\.|[-*])\s", lignes[i]):
                corps.append(lignes[i].strip()); i += 1
            puce(doc, inline(" ".join(corps)), num=marque); continue
        if not l.strip(): vider()
        else: para.append(l.strip())
        i += 1
    vider()
    if dst is None:
        dst = os.path.splitext(src)[0] + ".docx"
    doc.save(dst)
    return dst

if __name__ == "__main__":
    if len(sys.argv) < 2: sys.exit("usage: md2dossier.py <fichier.md> [sortie.docx]")
    print(convertir(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None))
