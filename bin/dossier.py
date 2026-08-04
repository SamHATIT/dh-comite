#!/usr/bin/env python3
"""Genere le dossier illustre du comite (daily ou hebdo) depuis l'etat en base.
Usage: dossier.py [daily|comite] [YYYY-MM-DD]"""
import os, sys, json, subprocess, datetime, re
import matplotlib; matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE="/workspace"; OUT=f"{BASE}/briefs"; TMP="/tmp/dh-charts"
INK=(0x0E,0x0E,0x10); BRASS=(0x8A,0x6D,0x3B); GREY=(0x6B,0x65,0x60)
CINK="#0E0E10"; CBRASS="#C8A97E"; CBRASS_D="#8A6D3B"; CGREY="#8D8880"
CVERT="#5E8C4E"; CAMBRE="#C79A3A"; CROUGE="#B0503F"; CPAPIER="#FBFAF7"
os.makedirs(TMP, exist_ok=True); os.makedirs(OUT, exist_ok=True)

plt.rcParams.update({"font.family":"DejaVu Sans","font.size":10,"text.color":CINK,
 "axes.edgecolor":"#D8D2C6","axes.labelcolor":CINK,"axes.facecolor":CPAPIER,
 "figure.facecolor":CPAPIER,"xtick.color":CGREY,"ytick.color":CGREY,
 "axes.grid":True,"grid.color":"#E6E1D7","grid.linewidth":.8})

def state(key):
    try:
        r=subprocess.run([f"{BASE}/bin/deos-state","get",key],capture_output=True,text=True,timeout=30)
        return json.loads(r.stdout) if r.stdout.strip() else None
    except Exception: return None

def decisions():
    dsn=os.environ.get("COMITE_DB_DSN","")
    sql=("SELECT json_agg(row_to_json(d)) FROM (SELECT id,statut,origine,left(texte,150) AS texte,"
         "(now()::date-date::date) AS age_j FROM decisions WHERE statut NOT IN ('clos','refusee') "
         "ORDER BY date) d;")
    try:
        r=subprocess.run(["psql",dsn,"-tA","-c",sql],capture_output=True,text=True,timeout=30)
        return json.loads(r.stdout) or [] if r.stdout.strip() else []
    except Exception: return []

def num(v, d=None):
    if isinstance(v,(int,float)): return v
    if isinstance(v,str):
        m=re.search(r"\d+",v)
        if m: return int(m.group())
    return d

def g_sante(_):
    lbl,val=[],[]
    noms={"delivery":"Delivery","commercial":"Commercial","marketing":"Marketing",
          "cs":"Customer Success","cos":"Execution (CoS)","legal":"Juridique"}
    for k,n in noms.items():
        r=state(f"rapport_{k}")
        s=num((r or {}).get("domain_score"))
        if s is not None: lbl.append(n); val.append(s)
    if not lbl: return None
    order=np.argsort(val); lbl=[lbl[i] for i in order]; val=[val[i] for i in order]
    col=[CVERT if v>=80 else (CAMBRE if v>=60 else CROUGE) for v in val]
    fig,ax=plt.subplots(figsize=(7.2,max(2.2,.62*len(lbl)+1.2)))
    b=ax.barh(lbl,val,color=col,height=.58,zorder=3)
    for r_,v in zip(b,val):
        ax.text(v+1.5,r_.get_y()+r_.get_height()/2,f"{v}",va="center",fontsize=11,fontweight="bold",color=CINK)
    ax.axvline(60,color=CAMBRE,ls="--",lw=1.2); ax.axvline(80,color=CVERT,ls="--",lw=1.2)
    ax.set_xlim(0,105); ax.set_xlabel("score du domaine /100"); ax.grid(axis="y",visible=False)
    ax.set_title("Sante par domaine",fontsize=11.5,color=CINK,pad=12,loc="left")
    p=f"{TMP}/sante.png"; fig.savefig(p,dpi=200,bbox_inches="tight",facecolor=CPAPIER); plt.close(fig)
    return p

def g_decisions(dec):
    if not dec: return None
    dec=sorted(dec,key=lambda d:-(d.get("age_j") or 0))[:14]
    lbl=[ (d["texte"][:42]+"...") if len(d["texte"])>42 else d["texte"] for d in dec ]
    age=[d.get("age_j") or 0 for d in dec]
    col=[CROUGE if a>14 else (CAMBRE if a>7 else CVERT) for a in age]
    fig,ax=plt.subplots(figsize=(7.2,max(2.4,.34*len(lbl)+1.3)))
    b=ax.barh(lbl[::-1],age[::-1],color=col[::-1],height=.62,zorder=3)
    for r_,v in zip(b,age[::-1]):
        ax.text(v+.35,r_.get_y()+r_.get_height()/2,f"{v} j",va="center",fontsize=9,color=CINK)
    ax.axvline(7,color=CAMBRE,ls="--",lw=1.2)
    ax.set_xlabel("jours d'attente"); ax.grid(axis="y",visible=False)
    ax.set_xlim(0,max(max(age)+3,10))
    ax.set_title("Decisions en attente d'arbitrage",fontsize=11.5,color=CINK,pad=12,loc="left")
    p=f"{TMP}/decisions.png"; fig.savefig(p,dpi=200,bbox_inches="tight",facecolor=CPAPIER); plt.close(fig)
    return p

def g_alertes(_):
    grav={"haute":0,"moyenne":0,"basse":0}
    for k in ["delivery","commercial","marketing","cs","cos"]:
        r=state(f"rapport_{k}") or {}
        for a in (r.get("alertes") or []):
            g=str(a.get("gravite","")).lower()
            for key in grav:
                if key in g: grav[key]+=1
    if sum(grav.values())==0: return None
    fig,ax=plt.subplots(figsize=(7.2,2.3))
    ks=["haute","moyenne","basse"]; vs=[grav[k] for k in ks]
    b=ax.barh(["Haute","Moyenne","Basse"],vs,color=[CROUGE,CAMBRE,CGREY],height=.55,zorder=3)
    for r_,v in zip(b,vs):
        ax.text(v+.08,r_.get_y()+r_.get_height()/2,str(v),va="center",fontsize=11,fontweight="bold")
    ax.set_xlim(0,max(vs)+1.4); ax.grid(axis="y",visible=False); ax.set_xlabel("nombre d'alertes")
    ax.set_title("Alertes par gravite, tous domaines",fontsize=11.5,color=CINK,pad=12,loc="left")
    p=f"{TMP}/alertes.png"; fig.savefig(p,dpi=200,bbox_inches="tight",facecolor=CPAPIER); plt.close(fig)
    return p

def g_fraicheur():
    lbl,h=[],[]
    noms={"delivery":"Delivery","commercial":"Commercial","marketing":"Marketing",
          "cs":"Customer Success","cos":"Execution (CoS)"}
    dsn=os.environ.get("COMITE_DB_DSN","")
    for k,n in noms.items():
        try:
            r=subprocess.run(["psql",dsn,"-tA","-c",
              f"SELECT round(extract(epoch FROM now()-updated_at)/3600,1) FROM deos_state WHERE cle='rapport_{k}';"],
              capture_output=True,text=True,timeout=20)
            v=float(r.stdout.strip()) if r.stdout.strip() else 999
        except Exception: v=999
        lbl.append(n); h.append(min(v,72))
    fig,ax=plt.subplots(figsize=(7.2,2.5))
    col=[CVERT if v<=24 else (CAMBRE if v<=48 else CROUGE) for v in h]
    b=ax.bar(lbl,h,color=col,width=.5,zorder=3)
    for r_,v in zip(b,h):
        ax.text(r_.get_x()+r_.get_width()/2,v+1,f"{v:.0f} h",ha="center",fontsize=9)
    ax.axhline(24,color=CAMBRE,ls="--",lw=1.2)
    ax.set_ylabel("anciennete du rapport (h)"); ax.grid(axis="x",visible=False)
    ax.set_title("Fraicheur des rapports - seuil d'alerte a 24 h",fontsize=11.5,color=CINK,pad=12,loc="left")
    plt.setp(ax.get_xticklabels(),fontsize=8.5)
    p=f"{TMP}/fraicheur.png"; fig.savefig(p,dpi=200,bbox_inches="tight",facecolor=CPAPIER); plt.close(fig)
    return p

def add(doc,txt,size=10.5,bold=False,italic=False,color=INK,space=6,align=None,font="Calibri"):
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(space)
    if align: p.alignment=align
    for part in re.split(r"(\*\*[^*]+\*\*)",str(txt)):
        if not part: continue
        b=part.startswith("**") and part.endswith("**")
        r=p.add_run(part[2:-2] if b else part)
        r.font.size=Pt(size); r.font.bold=b or bold; r.font.italic=italic
        r.font.name=font; r.font.color.rgb=RGBColor(*color)
    return p

def head(doc,txt,lvl):
    sizes={1:16,2:13,3:11.5}; colors={1:INK,2:BRASS,3:INK}
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(14 if lvl<3 else 10)
    p.paragraph_format.space_after=Pt(6)
    r=p.add_run(txt); r.font.size=Pt(sizes[lvl]); r.font.name="Georgia"
    r.font.color.rgb=RGBColor(*colors[lvl]); r.font.bold=(lvl==3); r.font.italic=(lvl==2)
    return p

def img(doc,path,width=16.2):
    if not path or not os.path.exists(path): return
    doc.add_picture(path,width=Cm(width))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER

def tbl(doc,headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""
        r=c.paragraphs[0].add_run(h); r.font.bold=True; r.font.size=Pt(9); r.font.name="Calibri"
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""
            for part in re.split(r"(\*\*[^*]+\*\*)",str(v)):
                if not part: continue
                b=part.startswith("**") and part.endswith("**")
                r=cells[i].paragraphs[0].add_run(part[2:-2] if b else part)
                r.font.size=Pt(9.5); r.font.bold=b; r.font.name="Calibri"
    return t

def main():
    kind=sys.argv[1] if len(sys.argv)>1 else "daily"
    day=sys.argv[2] if len(sys.argv)>2 else datetime.date.today().isoformat()
    brief=state("brief") or {}
    dec=decisions()
    doc=Document()
    st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(10.5)
    for s in doc.sections:
        s.top_margin=s.bottom_margin=Cm(1.9); s.left_margin=s.right_margin=Cm(1.9)
    titre="Comite de direction - edition hebdomadaire" if kind=="comite" else "Brief quotidien"
    add(doc,"DIGITAL\u00b7HUMANS",20,color=INK,space=2,align=WD_ALIGN_PARAGRAPH.CENTER,font="Georgia")
    add(doc,"C O M I T E   D E   D I R E C T I O N   A U G M E N T E",8,color=BRASS,space=18,
        align=WD_ALIGN_PARAGRAPH.CENTER)
    add(doc,titre,18,italic=True,space=4,align=WD_ALIGN_PARAGRAPH.CENTER,font="Georgia")
    add(doc,day,11,color=GREY,space=16,align=WD_ALIGN_PARAGRAPH.CENTER)
    sante=brief.get("sante") or {}
    score=num(sante.get("score"),"-")
    add(doc,f"Sante globale : {score}/100",14,bold=True,space=4,font="Georgia")
    if sante.get("tendance"): add(doc,f"Tendance : {sante['tendance']}",10,color=GREY,space=8)
    if sante.get("calcul"): add(doc,f"Calcul - {sante['calcul']}",9,italic=True,color=GREY,space=10)
    head(doc,"1. Vue d'ensemble",1)
    img(doc,g_sante(None)); img(doc,g_fraicheur())
    a=g_alertes(None)
    if a: img(doc,a)
    n=2
    if dec:
        head(doc,f"{n}. Decisions en attente",1); n+=1
        img(doc,g_decisions(dec))
        tbl(doc,["ID","Decision","Origine","Age"],
            [[d["id"],d["texte"][:110],d["origine"],f"{d.get('age_j',0)} j"] for d in dec[:14]])
    for key,title in [("hier","Hier - les faits"),("kpis","Indicateurs"),
              ("priorites_jour","Priorites"),("alertes","Alertes"),("opportunites","Opportunites")]:
        v=brief.get(key)
        if not v: continue
        head(doc,f"{n}. {title}",1); n+=1
        if isinstance(v,dict):
            for k2,v2 in v.items():
                add(doc,f"**{k2}** - {v2 if isinstance(v2,str) else json.dumps(v2,ensure_ascii=False)[:600]}",10)
        elif isinstance(v,list):
            for it in v[:12]:
                if isinstance(it,dict):
                    t=it.get("texte") or it.get("nom") or json.dumps(it,ensure_ascii=False)[:200]
                    extra=f" - {it['valeur']}" if it.get("valeur") else ""
                    src=f"  ({it['source']})" if it.get("source") else ""
                    add(doc,f"\u2022 {t}{extra}{src}",10,space=4)
                else: add(doc,f"\u2022 {it}",10,space=4)
        else: add(doc,str(v),10)
    reco=brief.get("recommandation")
    if reco:
        head(doc,f"{n}. Recommandation du CEO digital",1)
        add(doc,reco if isinstance(reco,str) else json.dumps(reco,ensure_ascii=False),10.5)
    add(doc,"",8,space=14)
    add(doc,"Document genere automatiquement par le comite - sources : deos_state, registre des decisions.",
        8,italic=True,color=GREY)
    path=f"{OUT}/{'comite' if kind=='comite' else 'brief'}-{day}.docx"
    doc.save(path); print(path)

if __name__=="__main__": main()
